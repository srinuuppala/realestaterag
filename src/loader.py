"""Corpus loading.

Reads every supported file in the knowledge base, extracts clean text, and
attaches the metadata that makes retrieval and citation useful: which builder,
which project, and what kind of document a chunk came from.
"""

from __future__ import annotations

import csv
import io
from collections.abc import Callable
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from src.config import get_settings
from src.utils import get_logger, humanise, normalise_whitespace, timed

log = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".html", ".htm", ".md", ".txt", ".csv"}

# The six flagship projects and three builders the corpus is built around.
PROJECT_CODES: dict[str, str] = {
    "sht": "Skyline Horizon Towers",
    "hbp": "Horizon Business Park",
    "mgr": "Meridian Greens Residency",
    "mlv": "Meridian Lakeview Villas",
    "unh": "Urban Nest Heights",
    "unr": "Urban Nest Riverside",
}

BUILDER_CODES: dict[str, str] = {
    "skyline": "Skyline Horizon Developers",
    "meridian": "Meridian Greens Realty",
    "urbannest": "Urban Nest Infrastructures",
    "propertybazaar": "PropertyBazaar (listing portal)",
}

PROJECT_TO_BUILDER: dict[str, str] = {
    "Skyline Horizon Towers": "Skyline Horizon Developers",
    "Horizon Business Park": "Skyline Horizon Developers",
    "Meridian Greens Residency": "Meridian Greens Realty",
    "Meridian Lakeview Villas": "Meridian Greens Realty",
    "Urban Nest Heights": "Urban Nest Infrastructures",
    "Urban Nest Riverside": "Urban Nest Infrastructures",
}

# Ordered longest-first so `terms_of_use` is not shadowed by `terms`.
CATEGORY_RULES: list[tuple[str, str]] = [
    ("cancellation_refund_policy", "Cancellation & Refund Policy"),
    ("possession_guidelines", "Possession Guidelines"),
    ("registration_process", "Registration Process"),
    ("sale_agreement_terms", "Sale Agreement / Legal Terms"),
    ("terms_conditions", "Terms & Conditions"),
    ("terms_of_use", "Terms & Conditions"),
    ("privacy_policy", "Privacy Policy"),
    ("customer_support", "Customer Support"),
    ("home_loan_partners", "Home Loan Information"),
    ("home_loan_general", "Home Loan Information"),
    ("builder_profile", "Builder Profile"),
    ("amenities_guide", "Amenities Guide"),
    ("location_guide", "Location Guide"),
    ("rera_general", "RERA Information"),
    ("rera_summary", "RERA Documentation"),
    ("payment_plan", "Payment Plan"),
    ("floor_plans", "Floor Plans"),
    ("brochure", "Project Brochure"),
    ("listing", "Property Listing"),
    ("about", "Builder Profile"),
    ("home", "Builder Website"),
    ("faq", "FAQ"),
]


# ---------------------------------------------------------------------------
# Format-specific text extraction
# ---------------------------------------------------------------------------
def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    document = DocxDocument(str(path))
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return "\n".join(blocks)


def _read_html(path: Path) -> str:
    soup = BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n")


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_csv(path: Path) -> str:
    """Rows become `column: value` lines so each row is semantically searchable."""
    raw = path.read_text(encoding="utf-8", errors="ignore")
    reader = csv.DictReader(io.StringIO(raw))
    lines: list[str] = []
    for index, row in enumerate(reader, start=1):
        pairs = [f"{key}: {value}" for key, value in row.items() if key and value]
        if pairs:
            lines.append(f"Row {index} — " + "; ".join(pairs))
    return "\n".join(lines)


READERS: dict[str, Callable[[Path], str]] = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".html": _read_html,
    ".htm": _read_html,
    ".md": _read_text,
    ".txt": _read_text,
    ".csv": _read_csv,
}


# ---------------------------------------------------------------------------
# Metadata
# ---------------------------------------------------------------------------
def classify(filename: str) -> dict[str, str]:
    """Derives project / builder / category from the file name.

    The corpus uses a strict `<code>_<category>.<ext>` convention, so this is
    reliable — and it is what lets a user ask "what is the payment plan for the
    Baner villas?" and have the retriever land on the right document.
    """
    stem = Path(filename).stem.lower()
    head = stem.split("_", 1)[0]

    project = PROJECT_CODES.get(head, "")
    builder = PROJECT_TO_BUILDER.get(project, BUILDER_CODES.get(head, ""))

    category = next((label for token, label in CATEGORY_RULES if token in stem), "General")

    return {
        "project": project or "All projects",
        "builder": builder or "General",
        "category": category,
    }


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def discover(root: Path | None = None) -> list[Path]:
    """Every supported file under the knowledge-base directory."""
    root = root or get_settings().knowledge_base_dir
    if not root.exists():
        raise FileNotFoundError(f"Knowledge base directory not found: {root}")

    return sorted(
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


@timed("load_documents")
def load_documents(root: Path | None = None) -> list[Document]:
    """Loads one Document per file. A file that fails to parse is skipped, not fatal."""
    documents: list[Document] = []

    for path in discover(root):
        reader = READERS[path.suffix.lower()]
        try:
            text = normalise_whitespace(reader(path))
        except Exception as error:  # noqa: BLE001 - one bad file must not sink the corpus
            log.error("Could not read %s: %s", path.name, error)
            continue

        if len(text) < 40:
            log.warning("Skipping %s — no extractable text", path.name)
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "source_path": str(path),
                    "title": humanise(path.name),
                    "file_type": path.suffix.lower().lstrip("."),
                    **classify(path.name),
                },
            )
        )

    log.info("Loaded %d documents", len(documents))
    return documents


@timed("split_documents")
def split_documents(documents: list[Document]) -> list[Document]:
    """Splits on natural boundaries (headings → paragraphs → sentences).

    Every chunk keeps its parent's metadata plus a heading breadcrumb, so a
    citation can say *which section* of a 12-page brochure was used.
    """
    settings = get_settings()
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        separators=["\n## ", "\n### ", "\n\n", "\n", ". ", " "],
        keep_separator=True,
    )

    chunks = splitter.split_documents(documents)

    for index, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = index
        chunk.metadata["section"] = _leading_heading(chunk.page_content)
        # Prepending the provenance line measurably improves embedding recall on
        # questions that name a project but not the document type.
        chunk.page_content = (
            f"[{chunk.metadata['project']} · {chunk.metadata['category']}]\n{chunk.page_content}"
        )

    log.info("Split into %d chunks", len(chunks))
    return chunks


def _leading_heading(text: str) -> str:
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            return stripped.lstrip("# ").strip()[:80]
    return ""


def build_corpus(root: Path | None = None) -> list[Document]:
    """Load → split. The single entry point used by the vector store."""
    return split_documents(load_documents(root))
